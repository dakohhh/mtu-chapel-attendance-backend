from docx import Document
from docx.shared import Pt
from docx.shared import Inches

def create_warning_letter(student_name, matric_no, reason):
    # Create a new Document
    doc = Document()

    # Add the university logo
    doc.add_picture('university_logo.jpeg', width=Inches(1.0))  # Adjust the path and width as needed
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = 1  # Center alignment

    # Add the text
    doc.add_paragraph('\n')  # Add some space after the logo
    p = doc.add_paragraph()
    p.add_run("This is to inform student ")
    p.add_run(student_name).bold = True
    p.add_run(", with matriculation number ")
    p.add_run(matric_no).bold = True
    p.add_run(" that he/she has been issued a warning letter for the reason ")
    p.add_run(reason).bold = True
    p.add_run(", and is hereby officially warned.")

    # Save the document
    doc.save('warning_letter.docx')

# Example usage
create_warning_letter("Ibiteye Marvellous", "20010301072", "ABSENCE IN CHAPEL")
