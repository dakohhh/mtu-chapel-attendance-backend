from docx import Document
from docx.shared import Inches



class WarningLetter():
    def __init__(self, student_name:str, matric_no:str, reason:str) -> None:

        self.student_name= student_name
        self.matric_no = matric_no
        self.reason = reason

    def create_warning_letter(self):
        # Create a new Document
        doc = Document()

        # Add the university logo
        doc.add_picture('university_logo.jpeg', width=Inches(1.0))  # Adjust the path and width as needed
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # Center alignment

        # Add the text
        doc.add_paragraph('\n')  # Add some space after the logo
        p = doc.add_paragraph()
        p.add_run("This is to inform student ")
        p.add_run(self.student_name).bold = True
        p.add_run(", with matriculation number ")
        p.add_run(str(self.matric_no)).bold = True
        p.add_run(" that he/she has been issued a warning letter for the reason ")
        p.add_run(self.reason).bold = True
        p.add_run(", and is hereby officially warned.")

        return doc
    